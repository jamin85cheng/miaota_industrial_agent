"""
文档加载器模块

功能需求: K-01 文档加载 - PDF/Word/Excel/Markdown
作者: Data Team + ML Team
"""

import os
from typing import List, Dict, Optional, Iterator
from pathlib import Path
from dataclasses import dataclass
from abc import ABC, abstractmethod
import tempfile
import shutil
from loguru import logger


@dataclass
class Document:
    """文档数据结构"""
    content: str
    metadata: Dict
    source: str
    doc_type: str
    page_number: Optional[int] = None


class BaseDocumentLoader(ABC):
    """文档加载器基类"""
    
    supported_extensions: List[str] = []
    
    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        """加载文档"""
        pass
    
    def can_load(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions


class PDFLoader(BaseDocumentLoader):
    """PDF 文档加载器"""
    
    supported_extensions = ['.pdf']
    
    def load(self, file_path: str, extract_images: bool = False) -> List[Document]:
        """
        加载 PDF 文档
        
        Args:
            file_path: PDF 文件路径
            extract_images: 是否提取图片中的文字
            
        Returns:
            文档列表（每页一个Document）
        """
        try:
            import PyPDF2
        except ImportError:
            logger.error("缺少PyPDF2依赖，请安装: pip install PyPDF2")
            return []
        
        documents = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    
                    if text.strip():
                        doc = Document(
                            content=text,
                            metadata={
                                'source': file_path,
                                'page': page_num,
                                'total_pages': len(pdf_reader.pages),
                                'title': pdf_reader.metadata.get('/Title', ''),
                                'author': pdf_reader.metadata.get('/Author', ''),
                            },
                            source=file_path,
                            doc_type='pdf',
                            page_number=page_num
                        )
                        documents.append(doc)
            
            logger.info(f"PDF加载完成: {file_path}, 共 {len(documents)} 页")
            
        except Exception as e:
            logger.error(f"PDF加载失败 {file_path}: {e}")
        
        return documents


class WordLoader(BaseDocumentLoader):
    """Word 文档加载器 (.docx)"""
    
    supported_extensions = ['.docx', '.doc']
    
    def load(self, file_path: str) -> List[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("缺少python-docx依赖，请安装: pip install python-docx")
            return []
        
        documents = []
        
        try:
            doc = DocxDocument(file_path)
            
            # 提取所有段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_rows.append(' | '.join(row_text))
                tables_text.append('\n'.join(table_rows))
            
            # 合并所有内容
            full_text = '\n\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n[表格内容]\n' + '\n\n'.join(tables_text)
            
            document = Document(
                content=full_text,
                metadata={
                    'source': file_path,
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables),
                },
                source=file_path,
                doc_type='word'
            )
            documents.append(document)
            
            logger.info(f"Word文档加载完成: {file_path}")
            
        except Exception as e:
            logger.error(f"Word文档加载失败 {file_path}: {e}")
        
        return documents


class ExcelLoader(BaseDocumentLoader):
    """Excel 文档加载器"""
    
    supported_extensions = ['.xlsx', '.xls', '.csv']
    
    def load(self, file_path: str, sheet_name: Optional[str] = None) -> List[Document]:
        """加载 Excel 文档"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("缺少pandas依赖")
            return []
        
        documents = []
        
        try:
            if file_path.endswith('.csv'):
                # CSV 文件
                df = pd.read_csv(file_path)
                content = self._dataframe_to_text(df)
                
                document = Document(
                    content=content,
                    metadata={
                        'source': file_path,
                        'rows': len(df),
                        'columns': len(df.columns),
                    },
                    source=file_path,
                    doc_type='csv'
                )
                documents.append(document)
            else:
                # Excel 文件
                xl_file = pd.ExcelFile(file_path)
                
                for sheet in xl_file.sheet_names:
                    if sheet_name and sheet != sheet_name:
                        continue
                    
                    df = pd.read_excel(file_path, sheet_name=sheet)
                    content = self._dataframe_to_text(df, sheet_name=sheet)
                    
                    document = Document(
                        content=content,
                        metadata={
                            'source': file_path,
                            'sheet': sheet,
                            'rows': len(df),
                            'columns': len(df.columns),
                        },
                        source=file_path,
                        doc_type='excel'
                    )
                    documents.append(document)
            
            logger.info(f"Excel加载完成: {file_path}, 共 {len(documents)} 个工作表")
            
        except Exception as e:
            logger.error(f"Excel加载失败 {file_path}: {e}")
        
        return documents
    
    def _dataframe_to_text(self, df, sheet_name: str = '') -> str:
        """将DataFrame转换为文本"""
        lines = []
        
        if sheet_name:
            lines.append(f"【工作表: {sheet_name}】")
        
        # 表头
        lines.append(' | '.join(df.columns.astype(str)))
        lines.append('-' * 50)
        
        # 数据行 (限制前100行)
        for _, row in df.head(100).iterrows():
            lines.append(' | '.join(str(v) for v in row.values))
        
        if len(df) > 100:
            lines.append(f'... ({len(df) - 100} 行省略)')
        
        return '\n'.join(lines)


class MarkdownLoader(BaseDocumentLoader):
    """Markdown 文档加载器"""
    
    supported_extensions = ['.md', '.markdown']
    
    def load(self, file_path: str) -> List[Document]:
        """加载 Markdown 文档"""
        documents = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 提取标题作为元数据
            import re
            title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
            title = title_match.group(1) if title_match else Path(file_path).stem
            
            document = Document(
                content=content,
                metadata={
                    'source': file_path,
                    'title': title,
                    'char_count': len(content),
                },
                source=file_path,
                doc_type='markdown'
            )
            documents.append(document)
            
            logger.info(f"Markdown加载完成: {file_path}")
            
        except Exception as e:
            logger.error(f"Markdown加载失败 {file_path}: {e}")
        
        return documents


class TextLoader(BaseDocumentLoader):
    """纯文本加载器"""
    
    supported_extensions = ['.txt', '.log', '.json', '.xml', '.py', '.js', '.java']
    
    def load(self, file_path: str) -> List[Document]:
        """加载文本文件"""
        documents = []
        
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"无法解码文件: {file_path}")
            
            document = Document(
                content=content,
                metadata={
                    'source': file_path,
                    'encoding': encoding,
                    'lines': content.count('\n') + 1,
                },
                source=file_path,
                doc_type='text'
            )
            documents.append(document)
            
            logger.info(f"文本文件加载完成: {file_path}")
            
        except Exception as e:
            logger.error(f"文本文件加载失败 {file_path}: {e}")
        
        return documents


class DocumentLoaderManager:
    """
    文档加载管理器
    
    统一管理所有文档加载器
    """
    
    def __init__(self):
        self.loaders: List[BaseDocumentLoader] = [
            PDFLoader(),
            WordLoader(),
            ExcelLoader(),
            MarkdownLoader(),
            TextLoader()
        ]
    
    def load(self, file_path: str) -> List[Document]:
        """
        加载文档（自动识别类型）
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档列表
        """
        file_path = str(file_path)
        
        for loader in self.loaders:
            if loader.can_load(file_path):
                return loader.load(file_path)
        
        logger.warning(f"不支持的文件类型: {file_path}")
        return []
    
    def load_directory(self, directory: str, 
                       recursive: bool = True,
                       extensions: Optional[List[str]] = None) -> Iterator[Document]:
        """
        加载目录中的所有文档
        
        Args:
            directory: 目录路径
            recursive: 是否递归子目录
            extensions: 只加载指定扩展名
            
        Yields:
            Document对象
        """
        directory = Path(directory)
        
        if not directory.exists():
            logger.error(f"目录不存在: {directory}")
            return
        
        pattern = '**/*' if recursive else '*'
        
        for file_path in directory.glob(pattern):
            if not file_path.is_file():
                continue
            
            # 过滤扩展名
            if extensions and file_path.suffix.lower() not in extensions:
                continue
            
            # 加载文档
            docs = self.load(str(file_path))
            for doc in docs:
                yield doc
    
    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型"""
        types = []
        for loader in self.loaders:
            types.extend(loader.supported_extensions)
        return types


# 使用示例
if __name__ == "__main__":
    # 创建管理器
    manager = DocumentLoaderManager()
    
    print("支持的文件类型:", manager.get_supported_types())
    
    # 测试加载不同文件
    test_files = [
        "data/knowledge_base/manual.pdf",
        "data/knowledge_base/guide.docx",
        "data/knowledge_base/data.xlsx",
        "README.md"
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\n加载: {file_path}")
            docs = manager.load(file_path)
            for doc in docs:
                print(f"  类型: {doc.doc_type}, 页数: {doc.page_number or 1}")
                print(f"  内容预览: {doc.content[:100]}...")
        else:
            print(f"文件不存在: {file_path}")
